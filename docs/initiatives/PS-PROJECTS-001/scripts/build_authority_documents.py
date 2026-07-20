"""Build Bible v2.6 and Roadmap v2.5 from the current controlled DOCX files.

This is intentionally a targeted, repeatable edit. It preserves the prior
versions as historical authority and changes only the Projects direction plus
the required version metadata.
"""

from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[4]
GOV = ROOT / "docs" / "governance"
BIBLE_SOURCE = GOV / "PeerSlate_Company_and_Product_Bible_v2.5.docx"
BIBLE_OUTPUT = GOV / "PeerSlate_Company_and_Product_Bible_v2.6.docx"
ROADMAP_SOURCE = GOV / "PeerSlate_Product_Strategy_and_Architecture_Roadmap_v2.4.docx"
ROADMAP_OUTPUT = GOV / "PeerSlate_Product_Strategy_and_Architecture_Roadmap_v2.5.docx"


def set_text(paragraph, text: str) -> None:
    """Replace visible text while preserving paragraph style and first-run format."""
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def paragraphs_in_tables(document):
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def all_paragraphs(document):
    yield from document.paragraphs
    yield from paragraphs_in_tables(document)


def replace_exact(document, old: str, new: str, expected: int = 1) -> None:
    matches = [p for p in all_paragraphs(document) if p.text == old]
    if len(matches) != expected:
        raise RuntimeError(f"Expected {expected} match(es) for {old!r}; found {len(matches)}")
    for paragraph in matches:
        set_text(paragraph, new)


def paragraph(document, exact: str):
    matches = [p for p in document.paragraphs if p.text == exact]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one body paragraph for {exact!r}; found {len(matches)}")
    return matches[0]


def table_containing(document, text: str):
    matches = [
        table
        for table in document.tables
        if any(text in cell.text for row in table.rows for cell in row.cells)
    ]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one table containing {text!r}; found {len(matches)}")
    return matches[0]


def table_with_header_containing(document, values: list[str], text: str):
    matches = [
        table
        for table in document.tables
        if table.rows
        and [cell.text for cell in table.rows[0].cells] == values
        and any(text in cell.text for row in table.rows for cell in row.cells)
    ]
    if len(matches) != 1:
        raise RuntimeError(
            f"Expected one table with header {values!r} containing {text!r}; "
            f"found {len(matches)}"
        )
    return matches[0]


def insert_before(anchor, entries: list[tuple[str, str]]) -> None:
    for style, text in entries:
        anchor.insert_paragraph_before(text, style=style)


def update_header_footer(
    document, header: str, footer_prefix: str, first_footer_text: str
) -> None:
    for section in document.sections:
        header_paragraph = section.header.paragraphs[0]
        set_text(header_paragraph, header)
        footer_paragraph = section.footer.paragraphs[0]
        if not footer_paragraph.runs:
            raise RuntimeError("Expected footer runs with a preserved PAGE field")
        footer_paragraph.runs[0].text = footer_prefix
        set_text(section.first_page_footer.paragraphs[0], first_footer_text)


def build_bible() -> None:
    doc = Document(BIBLE_SOURCE)
    replace_exact(doc, "v2.5 - Member-Directed Story Composition Authority", "v2.6 - Projects System Authority")
    replace_exact(doc, "July 18, 2026", "July 19, 2026")
    replace_exact(
        doc,
        "CURRENT - OWNER-APPROVED PRODUCT, VISUAL, AND STORY COMPOSITION AUTHORITY",
        "CURRENT - OWNER-APPROVED PRODUCT, VISUAL, STORY, AND PROJECTS AUTHORITY",
    )
    replace_exact(
        doc,
        "v2.4 visual-integrity authority + July 18 owner Member-Directed Story Composition decision",
        "v2.5 Story authority + July 19 owner Projects system decision",
    )
    replace_exact(
        doc,
        "Approval note: v2.5 is the current controlling Bible when named in CURRENT_BASELINE.yaml. It preserves the verified July 18 production and visual-integrity baseline and adds the owner-approved Member-Directed Story Composition covenant: members control Story arrangement, AI remains optional, layout metadata stays separate from canonical content, and layout save remains separate from publication.",
        "Approval note: v2.6 is the current controlling Bible when named in CURRENT_BASELINE.yaml. It preserves the verified production, visual-integrity, and member-directed Story baseline and adds the owner-approved Projects covenant: Projects are private-first connected containers, Project Workspaces and Project Projections remain separate, canonical records are linked rather than copied, and later implementation may not become a task-management clone.",
    )
    replace_exact(
        doc,
        "Authority: v2.5 is controlling when named by CURRENT_BASELINE.yaml. Prior Bible versions remain preserved as historical decision records.",
        "Authority: v2.6 is controlling when named by CURRENT_BASELINE.yaml. Prior Bible versions remain preserved as historical decision records.",
    )
    replace_exact(
        doc,
        "These startup and owner-understanding rules are active through the merged START_HERE pointer chain, templates, and governance guardrails. Visual-integrity and member-directed Story composition enforcement are active when v2.5, the Owner Visual Integrity Standard, and the Owner Story Composition Standard are named in CURRENT_BASELINE.yaml. The owner should not have to restate these procedures to a new agent or computer.",
        "These startup and owner-understanding rules are active through the merged START_HERE pointer chain, templates, and governance guardrails. Visual-integrity, member-directed Story composition, and Projects system enforcement are active when v2.6, the Owner Visual Integrity Standard, the Owner Story Composition Standard, and docs/initiatives/PS-PROJECTS-001/ are named in CURRENT_BASELINE.yaml. The owner should not have to restate these procedures to a new agent or computer.",
    )
    replace_exact(doc, "Project / Work item", "Project")
    replace_exact(
        doc,
        "A durable container connecting roles, contributions, decisions, outcomes, people, media, and related Moments.",
        "A private-first durable container connecting purpose, roles, contributions, decisions, milestones, outcomes, people, goals, skills, sources, media, and exact related Moment versions without copying authoritative content.",
    )

    boundary_anchor = paragraph(doc, "Not a conventional résumé builder, although it may create purpose-specific résumé exports from confirmed records.")
    boundary_anchor.insert_paragraph_before(
        "Not a project-management suite, task board, timesheet, procurement system, or enterprise delivery platform, although it may help a member remember, understand, and communicate project work.",
        style="List Bullet",
    )

    story_anchor = paragraph(doc, "Story as editorial curation")
    insert_before(
        story_anchor,
        [
            ("Heading 2", "Projects as durable member-owned containers"),
            (
                "Normal",
                "A Project is a private-first canonical container for a meaningful endeavor that unfolds across time. It connects exact approved records and relationships so the member can understand what moved, what they contributed, which decisions mattered, what happened, and what they may want to reuse. The Project does not copy the authoritative Moment, source, role, outcome, or member identity it connects.",
            ),
            (
                "List Bullet",
                "The private Project Workspace and any audience-specific Project Projection are separate objects and lifecycle states. Editing or completing the workspace never publishes it.",
            ),
            (
                "List Bullet",
                "Capture remains the common entry. A member confirms a Moment and then explicitly links its exact version to a Project through the governed Placement boundary.",
            ),
            (
                "List Bullet",
                "AI may propose titles, summaries, relationships, milestones, missing-context questions, retrospective prompts, or projection wording. Deterministic software and explicit member actions control creation, status, relationships, audience, publication, archive, and deletion.",
            ),
            (
                "List Bullet",
                "Work is the broader roles-and-contributions domain. A Project may sit within one role, span roles, or represent independent, learning, community, creative, or personal build work when the member chooses to connect it to professional growth.",
            ),
            (
                "List Bullet",
                "Slate Board Project notes are planning objects and may point to or propose a Project; they are not the canonical Project and cannot create one silently.",
            ),
        ],
    )

    ia_anchor = paragraph(doc, "14.3 Core functional behavior")
    ia_anchor.insert_paragraph_before(
        "PS-CORE-IA-012  PeerSlate shall distinguish the private Project Workspace, compact Work/My Slate Project summaries, and purpose- and audience-specific Project Projections without creating another top-level navigation forest or a second Project truth.",
        style="Requirement",
    )
    fr_anchor = paragraph(doc, "14.4 Data, provenance, and lifecycle")
    fr_anchor.insert_paragraph_before(
        "PS-CORE-FR-008  An authenticated member shall be able to create and manage a private Project, explicitly connect exact approved records, control Project lifecycle, and keep any audience projection as a separate reviewed and published state.",
        style="Requirement",
    )
    data_anchor = paragraph(doc, "14.5 AI quality and safety")
    data_anchor.insert_paragraph_before(
        "PS-CORE-DATA-006  Project records and relationships shall reference canonical records or exact governed versions without copying authoritative Capture, Moment, source, role, outcome, résumé, Story, or publication content into relationship state.",
        style="Requirement",
    )
    ai_anchor = paragraph(doc, "14.6 Security and privacy")
    ai_anchor.insert_paragraph_before(
        "PS-CORE-AI-006  AI may propose Project structure, relationships, questions, reflection, and projection wording, but shall not create, change status, complete, archive, delete, share, publish, or add collaborators without a separate explicit member action enforced by deterministic software.",
        style="Requirement",
    )
    sec_anchor = paragraph(doc, "14.7 Nonfunctional quality")
    sec_anchor.insert_paragraph_before(
        "PS-CORE-SEC-003  Project metadata, relationships, sources, media, workspace views, and projections shall be retrieved only after server authorization for owner, viewer, relationship, audience, purpose, lifecycle, and permitted source scope.",
        style="Requirement",
    )
    vnv_anchor = paragraph(doc, "14.8 Testing, verification, and validation")
    vnv_anchor.insert_paragraph_before(
        "PS-CORE-NFR-007  Every essential Project flow shall preserve readable semantic order and complete keyboard, touch, screen-reader, mobile, 200-percent-zoom, reduced-motion, long-content, missing-source, stale-edit, and failure-recovery behavior.",
        style="Requirement",
    )

    locked_anchor = paragraph(doc, "Member-directed Story composition is locked: the member controls supported item position, size, emphasis, and layering; AI remains an optional proposal; accessible alternatives and explicit publication are mandatory.")
    insert_before(
        locked_anchor,
        [
            ("Heading 2", "Projects vocabulary"),
            (
                "Normal",
                "Project — a private-first member-owned container connecting a meaningful endeavor to exact canonical records, relationships, lifecycle, and approved outcomes without copying their authoritative content.",
            ),
            (
                "Normal",
                "Project Workspace — the authenticated owner view for creating, linking, reviewing, correcting, organizing, completing, archiving, and deciding whether Project material should be reused or projected.",
            ),
            (
                "Normal",
                "Project Projection — a separately drafted, purpose- and audience-specific presentation of approved Project context with exact preview, explicit publication, traceability, and revocation.",
            ),
        ],
    )
    locked_anchor.insert_paragraph_before(
        "The Projects covenant is locked: Projects are private-first connected containers; workspace and projection remain separate; canonical records are linked, not copied; AI proposes; the member controls status, reuse, audience, and publication.",
        style="List Bullet",
    )
    later_anchor = paragraph(doc, "PS-STORY-COMPOSER-001 — Member-directed Story composition with owner-scoped layout revisions, accessible move/resize controls, responsive preview, and explicit publication after its future entry gate.")
    later_anchor.insert_paragraph_before(
        "PS-PROJECTS-001 — Private Projects Workspace, exact-version Moment links, Project lifecycle, connected reuse, and separately gated Project Projections after its future product, visual, architecture, and owner-route entry gate.",
        style="List Bullet",
    )

    doc.add_paragraph("Appendix L - Projects system covenant", style="Heading 1")
    doc.add_paragraph(
        "Projects are a natural later extension of the connected Slate. They preserve the history and meaning of an endeavor without turning PeerSlate into a task-management tool or reviving a fixture-driven portfolio island."
    )
    doc.add_paragraph("Constitutional acceptance rule", style="Heading 2")
    for text in (
        "A Project is private-first and owner-scoped. It connects exact approved canonical records and relationships rather than copying authoritative content.",
        "The private Project Workspace, compact summaries in connected domains, and any Project Projection remain separate purpose and lifecycle states.",
        "Project creation, status, relationships, audience, publication, archive, deletion, and collaboration remain member-controlled deterministic actions; AI may only propose.",
        "A future Project experience must be accessible, responsive, source-linked, honest about unavailable records, and useful even when never shared or published.",
        "Slate Board notes, historical public fixtures, résumé summaries, and demonstrations are not the authenticated canonical Projects product.",
    ):
        doc.add_paragraph(text, style="List Bullet")
    doc.add_paragraph("Current operational control", style="Heading 2")
    doc.add_paragraph(
        "The detailed requirements, experience direction, architecture, sequencing, and future entry gate are maintained in docs/initiatives/PS-PROJECTS-001/. The package is planned, not active, and changes no current route, schema, Project fixture, Slate Board behavior, or production capability."
    )

    update_header_footer(
        doc,
        "PeerSlate Company and Product Bible  |  Version 2.6 Projects system authority",
        "PEERSLATE • CURRENT BIBLE • JULY 19, 2026   |   ",
        "CURRENT AUTHORITY — VISUAL INTEGRITY, MEMBER-DIRECTED STORY, AND PROJECTS SYSTEM",
    )
    props = doc.core_properties
    props.title = "PeerSlate Company and Product Bible v2.6"
    props.subject = "Current product, visual integrity, Story composition, and Projects system authority"
    props.modified = datetime(2026, 7, 19)
    doc.save(BIBLE_OUTPUT)


def build_roadmap() -> None:
    doc = Document(ROADMAP_SOURCE)
    replace_exact(doc, "v2.4 - Member-Directed Story Composition Architecture", "v2.5 - Projects System Architecture")
    replace_exact(doc, "July 18, 2026", "July 19, 2026")
    replace_exact(
        doc,
        "CURRENT - OWNER-APPROVED SEQUENCE AND STORY COMPOSITION ARCHITECTURE",
        "CURRENT - OWNER-APPROVED SEQUENCE, STORY, AND PROJECTS ARCHITECTURE",
    )
    replace_exact(
        doc,
        "v2.3 clear-path roadmap + Bible v2.5 + July 18 Member-Directed Story Composition decision",
        "v2.4 Story roadmap + Bible v2.6 + July 19 Projects system decision",
    )
    replace_exact(
        doc,
        "Approval note: this v2.4 roadmap preserves the verified production and visual-integrity baseline and adds the owner-approved Story composition architecture: member-controlled move/resize/layering, accessible equivalents, separate versioned layout metadata, optional AI proposals, and explicit draft-versus-publication boundaries.",
        "Approval note: this v2.5 roadmap preserves the verified production, visual-integrity, and member-directed Story baseline and adds the owner-approved Projects architecture: a private Project Workspace, exact-version canonical links, member-controlled lifecycle, separately versioned Project Projections, optional AI proposals, and no task-management expansion.",
    )
    replace_exact(
        doc,
        "Roadmap authority: v2.4 is the current sequence and architecture when named by CURRENT_BASELINE.yaml. Historical package IDs and release evidence remain preserved; PS-STORY-COMPOSER-001 is planned future work, not active.",
        "Roadmap authority: v2.5 is the current sequence and architecture when named by CURRENT_BASELINE.yaml. Historical package IDs and release evidence remain preserved; PS-STORY-COMPOSER-001 and PS-PROJECTS-001 are planned future work, not active.",
    )
    replace_exact(
        doc,
        "CURRENT-STATE BASELINE / origin/main @ d5dd7bd…; pipeline 78 green; no open PRs; auth live; owner Settings live; private text Capture live; Deep Navy Gold live; Journal on hold. This snapshot must be superseded by repository CURRENT_STATE.md after every production-changing merge.",
        "CURRENT-STATE POINTER / CURRENT_BASELINE.yaml, CURRENT_STATE.md, and ACTIVE_INITIATIVES.md carry the verified repository, pipeline, active-lane, and production truth. Fetch origin before work; this Roadmap does not freeze a SHA or live-status snapshot.",
    )
    replace_exact(
        doc,
        "15. Phase 10 — Moment Lab, Story, Work, and connected views",
        "15. Phase 10 — Moment Lab, Story, Work, Projects, and connected views",
        expected=2,
    )
    replace_exact(
        doc,
        "Moment Lab, Story, Work, connected views",
        "Moment Lab, Story, Work, Projects, connected views",
    )
    replace_exact(
        doc,
        "Use the same canonical Slate to help members prepare for real professional moments and tell a curated story without duplicating identity or overwhelming the website.",
        "Use the same canonical Slate to help members prepare for real professional moments, keep Project work connected, and tell curated stories without duplicating identity or overwhelming the website.",
    )
    replace_exact(
        doc,
        "The current website already contains strong Interview Studio, My Story, Living Résumé, Slate Board, and project material. This phase does not add isolated replacements. It connects and simplifies them so one reviewed Moment can support preparation, Story, Work, projects, résumé exports, and selected public expression.",
        "The current website already contains strong Interview Studio, My Story, Living Résumé, Slate Board, and historical Project material. This phase does not revive isolated replacements. It creates the private Projects system and connects approved views so one reviewed Moment can support preparation, Story, Work, Projects, résumé exports, and selected public expression without fact copies.",
    )

    # There are many repeated headings; use the one immediately following the
    # Phase 10 architecture block by locating the nearest index after PS-WORK.
    body = doc.paragraphs
    phase_headings = [
        p
        for p in body
        if p.text == "15. Phase 10 — Moment Lab, Story, Work, Projects, and connected views"
        and p.style.name == "Heading 1"
    ]
    if len(phase_headings) != 1:
        raise RuntimeError(f"Expected one Phase 10 Heading 1; found {len(phase_headings)}")
    phase_heading = phase_headings[0]
    start_index = body.index(phase_heading)
    phase_anchor = next(p for p in body[start_index:] if p.text == "Minimum shall requirements")
    insert_before(
        phase_anchor,
        [
            ("Heading 2", "Projects Workspace and Project Projections"),
            (
                "Normal",
                "Projects receive a dedicated future package inside Phase 10. The first product is an authenticated private workspace that connects exact confirmed Moments and other governed records to one durable Project. A later Project Projection creates a purpose- and audience-specific case study or progress story only after exact preview and explicit publication.",
            ),
            (
                "List Bullet",
                "Do not revive the retired Pete-only Projects routes or promote historical fixtures into canonical member data without review, provenance, audience, and migration rules.",
            ),
            (
                "List Bullet",
                "Do not treat Slate Board Project notes, résumé Project summaries, or public cards as the system of record.",
            ),
            (
                "List Bullet",
                "The first slice is owner-only Project creation, lifecycle, exact-version Moment placement, and the accepted Project Ledger. Public projection, collaboration, task management, and homepage work remain separate later gates.",
            ),
        ],
    )

    work_req = paragraph(doc, "PS-WORK-001  Work and project views shall connect roles, contributions, decisions, outcomes, skills, sources, and related Moments without duplicate facts.")
    set_text(
        work_req,
        "PS-WORK-001  Work views shall connect roles, organizations, contributions, decisions, outcomes, skills, sources, and related Projects or Moments without duplicate facts.",
    )
    work_req.insert_paragraph_before(
        "PS-PROJECTS-001  The Projects system shall provide an owner-scoped private Project Workspace with explicit lifecycle, exact-version canonical links, progressive Project history, and no automatic publication.",
        style="Requirement",
    )
    work_req.insert_paragraph_before(
        "PS-PROJECTS-002  Project relationships shall reference exact governed records and shall not copy authoritative Capture, Moment, source, role, outcome, Story, résumé, or publication content.",
        style="Requirement",
    )
    work_req.insert_paragraph_before(
        "PS-PROJECTS-003  Project Workspace, connected-domain summaries, and Project Projections shall remain separate purpose, revision, audience, and publication states.",
        style="Requirement",
    )
    work_req.insert_paragraph_before(
        "PS-PROJECTS-004  AI may propose Project structure, relationships, reflection, and wording but shall not create, change lifecycle, share, publish, archive, delete, or add collaborators without explicit deterministic member action.",
        style="Requirement",
    )

    prototype_anchor = paragraph(doc, "Résumé/export review prototype with confirmed records, purpose-specific wording, source links, privacy check, and download/delete states.")
    prototype_anchor.insert_paragraph_before(
        "Project Workspace prototype with create, private state, empty/long ledgers, exact Moment links, lifecycle, stale edits, missing sources, archive/delete, mobile reflow, and one dominant action.",
        style="List Bullet",
    )
    prototype_anchor.insert_paragraph_before(
        "Later Project Projection prototype with purpose/audience selection, source-linked curation, exact responsive preview, private draft save, publish, revoke, and correction/deletion propagation.",
        style="List Bullet",
    )
    verification_anchor = paragraph(doc, "Work/project relationship and query tests; technical/sensitive context enforcement.")
    set_text(
        verification_anchor,
        "Work/Project relationship and query tests; two-owner isolation; exact-version Placement reuse; lifecycle and stale-write safety; no canonical-text copy; no automatic publication; technical/sensitive context enforcement.",
    )
    validation_anchor = paragraph(doc, "Each finds a Work/project contribution and reuses it in a purpose-specific draft without rewriting facts.")
    set_text(
        validation_anchor,
        "Each creates a private Project, links one exact confirmed Moment, explains Project versus Work/Slate Board/Projection, changes lifecycle safely, and reuses approved context in a purpose-specific draft without rewriting facts.",
    )

    phase_packages = table_with_header_containing(
        doc, ["ID", "Package", "Output"], "PS-MIA-MOMENTLAB-001"
    )
    for row in phase_packages.rows:
        if row.cells[0].text == "PS-WORK-001":
            set_text(row.cells[1].paragraphs[0], "Work and role views")
            set_text(row.cells[2].paragraphs[0], "Connected role/contribution/outcome history")
    phase_15a_matches = [
        p
        for p in doc.paragraphs
        if p.text
        == "15A. Future cross-phase initiative \u2014 Public Experience Visual Convergence"
        and p.style.name == "Heading 1"
    ]
    if len(phase_15a_matches) != 1:
        raise RuntimeError(
            f"Expected one Phase 15A Heading 1; found {len(phase_15a_matches)}"
        )
    phase_15a_anchor = phase_15a_matches[0]
    phase_15a_anchor.insert_paragraph_before(
        "Planned future package: PS-PROJECTS-001 \u2014 Projects Workspace and Project Projections \u2014 private Project foundation first; public projection and collaboration remain later gates.",
        style="Normal",
    )
    not_next_anchor = paragraph(doc, "Do not call work Complete because the happy-path page looks polished.")
    not_next_anchor.insert_paragraph_before(
        "Do not revive retired Projects pages, treat Slate Board notes as canonical Projects, or build a task-management suite from this direction package.",
        style="List Bullet",
    )

    package_folder_anchor = paragraph(doc, "Package folder standard")
    package_folder_anchor.insert_paragraph_before(
        "Planned governed extension: PS-PROJECTS-001 | Phase 10 | Projects Workspace and Project Projections | private owner Project, exact-version links, lifecycle, Project Ledger, and separately gated projections.",
        style="Normal",
    )

    doc.add_paragraph("Appendix F - Projects system direction amendment", style="Heading 1")
    doc.add_paragraph(
        "PS-PROJECTS-001 is planned Phase 10 expansion. It is not an active implementation lane and does not change current Project redirects, fixtures, Slate Board behavior, schema, or website capability."
    )
    doc.add_paragraph("Required sequence", style="Heading 2")
    for text in (
        "Validate the Project/Work/Slate Board boundaries and select one exact production-intent Project Workspace authority.",
        "Jointly baseline requirements and architecture for the owner-scoped Project aggregate, Slate entity registration, exact-version Moment placements, lifecycle, authorization, concurrency, migration, rollback, and data rights.",
        "Implement and release the private Project foundation and accepted Project Ledger before any public Project Projection or homepage promise.",
        "Add connected reuse through separately bounded Work, Story, Resume, Studio/Moment Lab, Replay, and export consumers.",
        "Authorize Project Projection and any collaboration only through later packages with exact audience, revision, publication, revocation, and access contracts.",
    ):
        doc.add_paragraph(text, style="List Bullet")
    doc.add_paragraph("Current control", style="Heading 2")
    doc.add_paragraph(
        "The controlling requirements, experience direction, architecture, traceability, test plan, exclusions, and entry gate are maintained in docs/initiatives/PS-PROJECTS-001/."
    )

    update_header_footer(
        doc,
        "PeerSlate Product Strategy and Architecture Roadmap  |  Version 2.5 Projects system architecture",
        "PEERSLATE • CURRENT ROADMAP • JULY 19, 2026   |   ",
        "CURRENT ROADMAP — SOURCE MATERIAL PRESERVED; PROJECTS SYSTEM DIRECTION ADOPTED",
    )
    props = doc.core_properties
    props.title = "PeerSlate Product Strategy and Architecture Roadmap v2.5"
    props.subject = "Current sequence and Projects system architecture"
    props.modified = datetime(2026, 7, 19)
    doc.save(ROADMAP_OUTPUT)


def main() -> None:
    build_bible()
    build_roadmap()
    print(BIBLE_OUTPUT)
    print(ROADMAP_OUTPUT)


if __name__ == "__main__":
    main()
